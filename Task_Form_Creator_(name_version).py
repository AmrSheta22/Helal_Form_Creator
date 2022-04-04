from PyQt5 import QtCore as qtc, QtGui as qtg, QtWidgets as qtw
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE
from PyQt5.QtGui import QPixmap
from pathlib import Path
import sys
import os
from docx2pdf import convert
import comtypes.client
class MainWindow(qtw.QWidget):
    def __init__(self):
        # importing the file and extracting ids
        
        df = pd.read_excel('target.xlsx')
        numbers = df['الاسم رباعي'].to_list()
        numbers = list(map(str, numbers))
        numbers.insert(0, "")
        df.set_index('الاسم رباعي', inplace= True)
        super().__init__()
        # Add a title                
        self.setStyleSheet("background-color: white;")
        self.setWindowTitle("استمارات الهلال الأحمر المصري")
        self.setWindowIcon(qtg.QIcon('helal.png'))
        #self.setFixedWidth(800)
        self.setFixedWidth(550)
        self.setMinimumHeight(600)        
        self.left = 500
        self.top = 200
        self.width= 1000
        self.height= 250
        # Set Vertical layout
        #self.setLayout(qtw.QVBoxLayout())
        form_layout = qtw.QFormLayout()
        group_box = qtw.QGroupBox()

        self.setLayout(form_layout)

        # Add Stuff/Widgets
        label_1 = qtw.QLabel("number of members")
        label_1.setAlignment(qtc.Qt.AlignCenter)
        label_1.setFont(qtg.QFont("Helvetica", 10))
        self.num_member = qtw.QLineEdit(self)
        self.num_member.setFixedWidth(300)
        self.num_member.setFixedHeight(30)
        self.num_member.setFont(qtg.QFont('Helvetica', 10))


        #image adding:
        self.label = qtw.QLabel(self)
        self.label.setAlignment(qtc.Qt.AlignCenter)

        self.pixmap = qtg.QPixmap(r'D:\Helal Project Files\helal.png')
 
        # adding image to label
        self.label.setPixmap(self.pixmap)
        form_layout.addRow(self.label)
        # Add Rows To App
        #form_layout.addRow(label_1)
        form_layout.addRow(label_1, self.num_member)
        

        self.bot = qtw.QPushButton("Add team",
            clicked = lambda: press_it(number_of_teams, chosen_members))
        self.bot.setFont(qtg.QFont('Helvetica', 11))
        self.bot.setFixedHeight(30)
        self.bot.setStyleSheet('QPushButton {background-color: red; color: white;}'
                                 'QPushButton::hover {background-color: green;}')
        form_layout.addRow(self.bot)
        
        number_of_teams = [1]
        chosen_members = []
        chosen_nums = []
        info_7 = []
        # Show the app

        self.show()
        
        def press_it(number_of_teams, chosen_members):

            label_2 = qtw.QLabel("                                          Team {}              ".format(number_of_teams[0]))
            label_2.setFont(qtg.QFont("Times", 10))

            form_layout.addRow(label_2)
            for i in range(int(self.num_member.text())):
                #add comboboxes of the number entered and modify it to be dropdown list
                
                label_3 = qtw.QLabel("Member {}".format(i+1))
                label_3.setFont(qtg.QFont("Helvetica", 9))
                globals()['my_combo{}'.format(i)] = qtw.QComboBox(self)
                globals()['my_combo{}'.format(i)].setEditable(True) 
                globals()['my_combo{}'.format(i)].completer().setCompletionMode(qtw.QCompleter.PopupCompletion)
                globals()['my_combo{}'.format(i)].setInsertPolicy(qtw.QComboBox.NoInsert)
                #adding items (the ids) to the list of the combobox
                globals()['my_combo{}'.format(i)].addItems(numbers)
                globals()['my_combo{}'.format(i)].activated
                globals()['my_combo{}'.format(i)].setFixedHeight(32)
                globals()['my_combo{}'.format(i)].setFont(qtg.QFont('Helvetica', 10))
                font = qtg.QFont('Helvetica', 10)
  
                line_edit = (globals()['my_combo{}'.format(i)]).lineEdit()
  
                line_edit.setFont(font)

                form_layout.addRow(label_3, globals()['my_combo{}'.format(i)])
            #creating bottuns and adding them to the layout
            self.newbot  = qtw.QPushButton("Finish Team Members", clicked = lambda: get_members_delete(chosen_members, chosen_nums))
            self.newbot.setFont(qtg.QFont('Helvetica', 11))
            self.newbot.setFixedHeight(30)
            self.newbot.setStyleSheet('QPushButton {background-color: red; color: white;}'
                                     'QPushButton::hover {background-color: green;}')
            self.newbot2  = qtw.QPushButton("Next Team", clicked = lambda: create_team())
            self.newbot2.setFont(qtg.QFont('Helvetica', 11))
            self.newbot2.setFixedHeight(30)
            self.newbot2.setStyleSheet('QPushButton {background-color: red; color: white;}'
                                      'QPushButton::hover {background-color: green;}')
            form_layout.addRow(self.newbot)
            form_layout.addRow(self.newbot2)
            number_of_teams[0]+=1
            #getting the ids and adding the names of the ids to the list
        def check_boxes():
            check = True
            for i in range(int(self.num_member.text())):
                if globals()['my_combo{}'.format(i)].currentText() == "":
                    check= False
            return check
                    
        def get_members_delete(chosen_members, chosen_nums):
            if check_boxes():
                get_members(chosen_members, chosen_nums)
                for i in range(3, 8+int(self.num_member.text())):
                    form_layout.removeRow(1)
                finish_first_app()
        def finish_first_app():

            
            self.text2 = qtw.QLineEdit(self)
            self.text2.setFixedWidth(300)
            self.text2.setFixedHeight(30)
            self.text2.setFont(qtg.QFont('Helvetica', 10))

            self.text3 = qtw.QLineEdit(self)
            self.text3.setFixedWidth(300)
            self.text3.setFixedHeight(30)
            self.text3.setFont(qtg.QFont('Helvetica', 10))

            self.text4 = qtw.QLineEdit(self)
            self.text4.setFixedWidth(300)
            self.text4.setFixedHeight(30)
            self.text4.setFont(qtg.QFont('Helvetica', 10))

            self.text5 = qtw.QLineEdit(self)
            self.text5.setFixedWidth(300)
            self.text5.setFixedHeight(30)
            self.text5.setFont(qtg.QFont('Helvetica', 10))

            self.text6 = qtw.QLineEdit(self)
            self.text6.setFixedWidth(300)
            self.text6.setFixedHeight(30)
            self.text6.setFont(qtg.QFont('Helvetica', 10))

            self.text7 = qtw.QLineEdit(self)
            self.text7.setFixedWidth(300)
            self.text7.setFixedHeight(30)
            self.text7.setFont(qtg.QFont('Helvetica', 10))



            label_12 = qtw.QLabel("Task Name")
            label_12.setAlignment(qtc.Qt.AlignCenter)
            label_12.setFont(qtg.QFont("Helvetica", 10))

            label_13 = qtw.QLabel("Task Date")
            label_13.setAlignment(qtc.Qt.AlignCenter)
            label_13.setFont(qtg.QFont("Helvetica", 10))

            label_14 = qtw.QLabel("Leader Name")
            label_14.setAlignment(qtc.Qt.AlignCenter)
            label_14.setFont(qtg.QFont("Helvetica", 10))

            label_15 = qtw.QLabel("Points")
            label_15.setAlignment(qtc.Qt.AlignCenter)
            label_15.setFont(qtg.QFont("Helvetica", 10))

            label_16 = qtw.QLabel("Task Start Time")
            label_16.setAlignment(qtc.Qt.AlignCenter)
            label_16.setFont(qtg.QFont("Helvetica", 10))

            label_17 = qtw.QLabel("Task End Time")
            label_17.setAlignment(qtc.Qt.AlignCenter)
            label_17.setFont(qtg.QFont("Helvetica", 10))

            form_layout.addRow(label_12, self.text2)
            form_layout.addRow(label_13, self.text3)
            form_layout.addRow(label_14, self.text4)
            form_layout.addRow(label_15, self.text5)
            form_layout.addRow(label_16, self.text6)
            form_layout.addRow(label_17, self.text7)

            self.newbot11  = qtw.QPushButton("finish", clicked = lambda: get_info(info_7))
            self.newbot11.setFont(qtg.QFont('Helvetica', 11))
            self.newbot11.setFixedHeight(30)
            self.newbot11.setStyleSheet('QPushButton {background-color: red; color: white;}'
                                     'QPushButton::hover {background-color: green;}')
            form_layout.addRow(self.newbot11)
            self.setFixedHeight(400)        

        def get_info(info_7):

            info_7.append(self.text2.text())
            info_7.append(self.text3.text())
            info_7.append(self.text4.text())
            info_7.append(self.text5.text())
            info_7.append(self.text6.text())
            info_7.append(self.text7.text())
            file_number = 1
            while Path('ALEX-{}-{}.pdf'.format(file_number, info_7[1])).is_file():
                file_number+=1
            if file_number>10:
                info_7.insert(0, 'ALEX-{}-{}'.format(file_number, info_7[1]))
            info_7.insert(0, 'ALEX-0{}-{}'.format(file_number, info_7[1]))
 
            finish_app(info_7,chosen_members, chosen_nums)        

        def finish_app(info_7,chosen_members, chosen_nums):

            document = Document("form.docx")
            add_task_info(document, info_7[0], info_7[1],info_7[2],info_7[3], info_7[4],info_7[5], info_7[6])
            for i, (mem, num) in enumerate(zip(chosen_members, chosen_nums)):
                title = document.add_heading(" {} الفريق رقم".format(i+1), 1)
                p = document.add_paragraph("")
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_style = title.style
                title_style.font.name = "Times New Roman"
                title_style.font.size = docx.shared.Pt(14)
                title_style.font.bold = True
                if len(mem)%2==1:
                    mem.append("")
                    num.append("")
                create_doc(document, num, mem)
                p = document.add_paragraph("")

            while True:
                if "/" in info_7[0]:
                    info_7[0]= info_7[0].replace(info_7[0][info_7[0].index("/")], "-", 1)
                else:
                    break
            document.save('{}.docx'.format(info_7[0]))
            convert('{}.docx'.format(info_7[0]))
            convert('{}.docx'.format(info_7[0]), '{}.pdf'.format(info_7[0]))
            os.remove('{}.docx'.format(info_7[0]))

            self.close()



        def get_members(chosen_members, chosen_nums):
            temp_list = []
            temp_num = []
            for i in range(int(self.num_member.text())):
                if globals()['my_combo{}'.format(i)].currentText() != "":
                    temp_list.append(int(df.loc[globals()['my_combo{}'.format(i)].currentText(), ['رقم العضوية لعام 2021']].values[0]))
                    temp_num.append(globals()['my_combo{}'.format(i)].currentText())
            chosen_nums.append(temp_num)
            chosen_members.append(temp_list)
        #creating another team option
        def create_team():
            get_members(chosen_members,chosen_nums)
            for i in range(3, 6+int(self.num_member.text())):
                form_layout.removeRow(3)
            self.num_member.clear()
        
        ##############################################################################
        def add_task_info(document, task_code, task_name , task_date, task_time, points, leader_name, member_count):
            table1 = document.add_table(rows=7, cols=2, style= 'Grid Table 2 Accent 2')

            hdr_cells = table1.columns[1].cells
            hdr_cells[0].text = 'كود المهمة'
            hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[0].paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
            hdr_cells[0].paragraphs[0].runs[0].font.bold = True

            hdr_cells[1].text = 'اسم المهمة'
            hdr_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[1].paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
            hdr_cells[1].paragraphs[0].runs[0].font.bold = True

            hdr_cells[2].text = 'تاريخ المهمة'
            hdr_cells[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[2].paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
            hdr_cells[2].paragraphs[0].runs[0].font.bold = True

            hdr_cells[3].text = 'اسم مسئول المهمة'
            hdr_cells[3].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[3].paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
            hdr_cells[3].paragraphs[0].runs[0].font.bold = True

            hdr_cells[4].text = 'عدد النقاط'
            hdr_cells[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[4].paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
            hdr_cells[4].paragraphs[0].runs[0].font.bold = True

            hdr_cells[5].text = 'توقيت بدأ المهمة'
            hdr_cells[5].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[5].paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
            hdr_cells[5].paragraphs[0].runs[0].font.bold = True

            hdr_cells[6].text = 'توقيت نهاية المهمة'
            hdr_cells[6].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[6].paragraphs[0].runs[0].font.size = docx.shared.Pt(14)
            hdr_cells[6].paragraphs[0].runs[0].font.bold = True

            ######################################
            hdr_cells = table1.columns[0].cells
            hdr_cells[0].text = task_code
            hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[0].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)

            hdr_cells[1].text = task_name
            hdr_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[1].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)

            hdr_cells[2].text = task_date
            hdr_cells[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[2].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)

            hdr_cells[3].text = task_time
            hdr_cells[3].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[3].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)

            hdr_cells[4].text = points
            hdr_cells[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[4].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)

            hdr_cells[5].text = leader_name
            hdr_cells[5].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[5].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)

            hdr_cells[6].text = member_count
            hdr_cells[6].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[6].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)

            for i in range(7):
                table1.cell(i,1).width = Inches(2.6)
                table1.cell(i,0).width = Inches(5)
            document.add_paragraph("")
                
        def create_doc(document, names, numbers):
            document.add_paragraph("")
            table = document.add_table(rows=1, cols=6, style= 'Grid Table 4 Accent 2')

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'رقم العضوية'
            hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[0].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)

            hdr_cells[1].text = 'أسماء المتطوعين'
            hdr_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[1].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)

            hdr_cells[3].text = 'رقم العضوية'
            hdr_cells[3].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[3].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)

            hdr_cells[4].text = 'أسماء المتطوعين'
            hdr_cells[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[4].paragraphs[0].runs[0].font.size = docx.shared.Pt(13)


            for i in range(0,len(numbers), 2):

                row_cells = table.add_row().cells

                row_cells[0].text = str(numbers[i+1])
                row_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row_cells[0].paragraphs[0].runs[0].font.size = docx.shared.Pt(12)
                row_cells[0].paragraphs[0].runs[0].font.bold = True

                row_cells[1].text = names[i+1]
                row_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row_cells[1].paragraphs[0].runs[0].font.size = docx.shared.Pt(12)
                row_cells[1].paragraphs[0].runs[0].font.bold = True


                row_cells[3].text = str(numbers[i])
                row_cells[3].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row_cells[3].paragraphs[0].runs[0].font.size = docx.shared.Pt(12)
                row_cells[3].paragraphs[0].runs[0].font.bold = True


                row_cells[4].text = names[i]
                row_cells[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row_cells[4].paragraphs[0].runs[0].font.size = docx.shared.Pt(12)
                row_cells[4].paragraphs[0].runs[0].font.bold = True


            count_rows =0
            for row in table.rows:
                count_rows+=1
                row.height_rule = WD_ROW_HEIGHT.EXACTLY
                row.height = docx.shared.Pt(20)
            table.rows[0].height = docx.shared.Pt(24)
            for i in range(count_rows):
                table.cell(i,0).width = Inches(1.0)
                table.cell(i,3).width = Inches(1.0)
                if i < count_rows-1:
                    table.cell(i+1,2).text = str(count_rows+(i+1)-1)
                    table.cell(i+1,5).text = str(i+1)
                table.cell(i,2).width = Inches(.3)
                table.cell(i,5).width = Inches(.3)
                table.cell(i,1).width = Inches(2.8)
                table.cell(i,4).width = Inches(2.8)
        ############################################################################
        form_layout.setAlignment(qtc.Qt.AlignCenter)

        form_layout.getWidgetPosition

        group_box.setLayout(form_layout)
        group_box.setFixedWidth(500)
        

        scroll = qtw.QScrollArea()
        scroll.setWidget(group_box)
        scroll.setWidgetResizable(True)

        layout = qtw.QVBoxLayout()
        layout.addWidget(scroll)
        self.setLayout(layout)
            
app = qtw.QApplication([])
mw = MainWindow()
app.exec_()